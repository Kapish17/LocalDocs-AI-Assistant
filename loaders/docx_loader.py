from langchain_core.documents import Document
import docx


def load_docx(file_path):
    """
    Reads a DOCX file and extracts text.
    """

    document = docx.Document(file_path)

    text = ""

    for paragraph in document.paragraphs:
        text += paragraph.text + "\n"

    return Document(
        page_content=text,
        metadata={
            "source": file_path.name,
            "file_type": "docx"
        }
    )